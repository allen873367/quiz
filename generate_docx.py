# -*- coding: utf-8 -*-
"""生成專題計畫書 Word 檔"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── 全域字型設定 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

# 設定邊界
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


def set_cell_shading(cell, color):
    """設定表格儲存格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(text, level=1):
    """加入標題並設定中文字型"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    return h


def add_para(text, bold=False, align=None, size=None, space_after=Pt(6)):
    """加入段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    run.bold = bold
    if size:
        run.font.size = size
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    return p


def add_table(headers, rows, col_widths=None):
    """加入表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表頭
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2B579A')
        # 白色字
        run.font.color.rgb = RGBColor(255, 255, 255)

    # 資料行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
            run.font.size = Pt(11)
            if r_idx % 2 == 0:
                set_cell_shading(cell, 'F2F7FB')

    # 欄寬
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # 表後空行
    return table


# ══════════════════════════════════════════
# 封面
# ══════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

add_para('資料結構', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(28))
add_para('主題測驗網站', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(28))

doc.add_paragraph()

add_para('專題計畫書', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(22))

for _ in range(4):
    doc.add_paragraph()

add_para('授課教師：林明言 老師', align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(14))
doc.add_paragraph()
add_para('CBF113007 鄭書懷', align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(14))
add_para('CBF113018 郭漢廷', align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(14))

doc.add_paragraph()
add_para('中華民國 114 年 6 月', align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(14))

doc.add_page_break()

# ══════════════════════════════════════════
# 目錄（手動）
# ══════════════════════════════════════════
add_heading_styled('目錄', level=1)
toc_items = [
    '一、專題名稱',
    '二、團隊成員',
    '三、專案緣起與動機',
    '四、專案目標',
    '五、系統架構',
    '六、功能規格',
    '七、連貫題機制說明',
    '八、間隔學習演算法',
    '九、開發時程',
    '十、資料庫操作說明',
    '十一、URL 路由一覽',
    '十二、未來展望',
    '十三、參考資料',
]
for item in toc_items:
    add_para(item, size=Pt(12))

doc.add_page_break()

# ══════════════════════════════════════════
# 一、專題名稱
# ══════════════════════════════════════════
add_heading_styled('一、專題名稱', level=1)
add_para('資料結構主題測驗網站（Data Structure Quiz Website）', bold=True, size=Pt(14))

# ══════════════════════════════════════════
# 二、團隊成員
# ══════════════════════════════════════════
add_heading_styled('二、團隊成員', level=1)
add_table(
    ['學號', '姓名', '職稱', '主要負責項目'],
    [
        ['CBF113007', '鄭書懷', '組員', '後端開發、測驗邏輯、資料庫設計'],
        ['CBF113018', '郭漢廷', '組員', '前端設計、主題樣式、UI/UX'],
    ],
    col_widths=[3, 2.5, 2, 7]
)

# ══════════════════════════════════════════
# 三、專案緣起與動機
# ══════════════════════════════════════════
add_heading_styled('三、專案緣起與動機', level=1)

add_para('在大專院校的資料結構課程中，「樹狀結構」（Tree）是核心章節之一，包含二元樹、完滿二元樹、完整二元樹、二元搜尋樹等重要概念。然而傳統的紙本測驗存在以下限制：')

problems = [
    ('缺乏即時回饋', '學生需等待老師批改後才知道對錯，無法立即得知學習成果。'),
    ('無法針對弱點強化', '一次測驗結束後，錯題容易被遺忘，缺乏有效的複習機制。'),
    ('學習歷程難以追蹤', '無法系統性回顧過往作答情況，進步與弱點不易量化。'),
    ('出題彈性低', '每次測驗需手動選題，無法隨機組合或自訂題數。'),
]

for title, desc in problems:
    p = doc.add_paragraph()
    run_b = p.add_run(f'• {title}：')
    run_b.bold = True
    run_b.font.name = 'Times New Roman'
    run_b.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    run_n = p.add_run(desc)
    run_n.font.name = 'Times New Roman'
    run_n.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

add_para('')
add_para('為了解決上述問題，本專題開發一個以樹狀結構為主題的線上測驗網站，運用間隔學習（Spaced Repetition）演算法幫助學生鞏固弱點，並提供完整的學習歷程記錄與排行系統，以提升學習動機與效果。')

# ══════════════════════════════════════════
# 四、專案目標
# ══════════════════════════════════════════
add_heading_styled('四、專案目標', level=1)

add_heading_styled('4.1 核心目標', level=2)

goals = [
    ('建立完整的樹狀結構題庫', '收錄 23 題涵蓋二元樹、完滿二元樹、完整二元樹、二元搜尋樹、樹的追蹤（前序／中序／後序）等核心概念。'),
    ('實現彈性的測驗機制', '支援章節測驗、隨機挑戰、自訂題數（5/10/15/20/全部）。'),
    ('導入間隔學習演算法', '錯題自動進入下一輪，直到全部答對，強化長期記憶。'),
    ('提供學習歷程記錄', '完整保存每次測驗成績、錯題、作答時間。'),
    ('建立排行系統', '透過一般排行與間隔學習排行，提升學習動機。'),
]

for title, desc in goals:
    p = doc.add_paragraph()
    run_b = p.add_run(f'{title}：')
    run_b.bold = True
    run_b.font.name = 'Times New Roman'
    run_b.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    run_n = p.add_run(desc)
    run_n.font.name = 'Times New Roman'
    run_n.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

add_heading_styled('4.2 延伸目標', level=2)
ext_goals = [
    '自編題目擴充題庫，增加題目多樣性。',
    '隨機挑戰支援跨章節混合出題，不受單一章節限制。',
    '連貫題組保護機制，確保同組題目（如 11→12→13）整組出現不被拆散。',
    '管理後台支援題目、用戶、記錄的完整 CRUD 操作。',
]
for g in ext_goals:
    add_para(f'• {g}')

# ══════════════════════════════════════════
# 五、系統架構
# ══════════════════════════════════════════
add_heading_styled('五、系統架構', level=1)

add_heading_styled('5.1 整體架構', level=2)
add_para('本系統採用 Django MTV（Model-Template-View）架構，前端以 HTML/CSS/JavaScript 渲染，後端以 Python Django 處理商務邏輯，資料庫使用 SQLite。使用者透過瀏覽器操作，所有互動皆透過 HTTP 請求與伺服器溝通。')

add_para('')
add_para('系統運作流程：', bold=True)
steps = [
    '使用者透過瀏覽器訪問網站，送出 HTTP 請求。',
    'Django URL 路由（urls.py）將請求分派至對應的檢視函數（views.py）。',
    '檢視函數從資料庫（models.py）讀取或寫入資料。',
    '檢視函數將資料傳遞給模板（templates），渲染成 HTML 回傳給瀏覽器。',
    '前端 CSS（Aurora Glass 主題）負責視覺呈現與動畫效果。',
]
for i, s in enumerate(steps, 1):
    add_para(f'{i}. {s}')

add_heading_styled('5.2 資料模型設計', level=2)
add_para('本系統包含四個主要資料模型，其關係如下：')
add_para('')

# 用文字說明取代圖
add_para('User（使用者）', bold=True)
add_para('    ├── 暱稱（nickname）、帳號（username）、Email、密碼、管理員權限')
add_para('    │')
add_para('    ├── QuizRecord（測驗記錄）— 一對多', bold=True)
add_para('    │    ├── 章節、總題數、答對數、分數、耗時、是否 SR、作答時間')
add_para('    │    │')
add_para('    │    └── WrongAnswer（錯題記錄）— 一對多', bold=True)
add_para('    │         ├── 對應題目（Question）、使用者答案、作答時間')
add_para('    │')
add_para('    └── Question（題目）— 多對多（透過 WrongAnswer）', bold=True)
add_para('         ├── 章節、題號、題目內容、題目圖片（可選）')
add_para('         ├── 選項 A~E（D、E 可選留空）')
add_para('         ├── 正確答案、難易度、題目詳解')
add_para('         ├── 連貫題組別（sequence_group）、是否自編題、來源說明')
add_para('         └── 選項作答時隨機排列，正確答案動態對應')

add_para('')
add_heading_styled('5.3 技術棧', level=2)
add_table(
    ['層級', '技術', '版本'],
    [
        ['後端框架', 'Django', '6.0.5'],
        ['程式語言', 'Python', '3.14'],
        ['資料庫', 'SQLite', '—'],
        ['前端樣式', 'CSS（Aurora Glass 主題）', '自訂'],
        ['圖片處理', 'Pillow', '—'],
        ['題目匯入', 'python-docx', '1.2.0'],
    ],
    col_widths=[3, 6, 3]
)

# ══════════════════════════════════════════
# 六、功能規格
# ══════════════════════════════════════════
add_heading_styled('六、功能規格', level=1)

add_heading_styled('6.1 用戶系統', level=2)
add_table(
    ['功能', '說明', '優先級'],
    [
        ['註冊', '填寫帳號、暱稱、密碼即可註冊新帳號', '★★★'],
        ['登入／登出', '表單認證登入與 Session 管理', '★★★'],
        ['個人資料編輯', '修改暱稱、Email', '★★'],
        ['密碼更改', '需提供舊密碼驗證後方可更改', '★★'],
        ['帳號刪除', '輸入 DELETE 確認後永久刪除帳號', '★★'],
    ],
    col_widths=[3, 8, 2]
)

add_heading_styled('6.2 測驗系統', level=2)
add_table(
    ['功能', '說明', '優先級'],
    [
        ['章節測驗', '針對第 7 章樹狀結構出題', '★★★'],
        ['隨機挑戰', '跨章節混合出題，自由選擇範圍', '★★★'],
        ['自訂題數', '5/10/15/20/全部，彈性選擇', '★★★'],
        ['連貫題保護', '同組題目（11→12→13）整組出現不拆散', '★★★'],
        ['即時回饋', '每題作答後立即顯示正確/錯誤與詳細解析', '★★★'],
        ['計時器', '測驗過程即時顯示已耗時間', '★★'],
        ['選項隨機', '每題選項內容隨機排列，防止背答案', '★★'],
    ],
    col_widths=[3, 8, 2]
)

add_heading_styled('6.3 間隔學習（Spaced Repetition）', level=2)
add_table(
    ['功能', '說明', '優先級'],
    [
        ['第一輪評分', '以第一輪答對數作為最終評分依據', '★★★'],
        ['錯題重複', '答錯的題目自動進入下一輪', '★★★'],
        ['答對移除', '答對的題目在下一輪不再出現', '★★★'],
        ['無上限輪次', '持續進行直到所有題目答對為止', '★★★'],
        ['SR 排行', '依總完成時間排名（時間越短越前面）', '★★'],
    ],
    col_widths=[3, 8, 2]
)

add_heading_styled('6.4 學習記錄與排行榜', level=2)
add_table(
    ['功能', '說明', '優先級'],
    [
        ['作答記錄', '自動保存每次測驗的成績、時間、章節', '★★★'],
        ['錯題列表', '按測驗記錄分組顯示錯題', '★★★'],
        ['錯題詳情', '查看題目、選項、正確答案與完整解析', '★★★'],
        ['一般排行', '依分數排名，顯示前 20 名', '★★★'],
        ['SR 排行', '依完成時間排名，顯示前 20 名', '★★'],
        ['排行篩選', '支援依章節、題數過濾排行結果', '★★'],
    ],
    col_widths=[3, 8, 2]
)

add_heading_styled('6.5 管理後台', level=2)
add_table(
    ['功能', '說明', '優先級'],
    [
        ['題目 CRUD', '新增、編輯、刪除題目（含圖片上傳）', '★★★'],
        ['用戶管理', '編輯權限、新增、刪除用戶', '★★★'],
        ['記錄管理', '檢視與刪除測驗記錄', '★★'],
        ['批次刪除', '全選後批次刪除題目/用戶/記錄', '★★'],
        ['搜尋篩選', '依章節、難易度、題號搜尋題目', '★★'],
        ['系統統計', '題庫總數、自編題數、用戶數、測驗次數總覽', '★'],
    ],
    col_widths=[3, 8, 2]
)

add_heading_styled('6.6 UI/UX', level=2)
add_table(
    ['功能', '說明', '優先級'],
    [
        ['Aurora Glass 主題', '深紫藍漸層背景、玻璃擬態卡片、圓角按鈕', '★★★'],
        ['首頁動畫', '流星雨背景效果、滑鼠游標極光光暈', '★★'],
        ['深色主題表單', '管理後台下拉選單統一深色風格', '★★'],
    ],
    col_widths=[3, 8, 2]
)

# ══════════════════════════════════════════
# 七、連貫題機制說明
# ══════════════════════════════════════════
add_heading_styled('七、連貫題機制說明', level=1)

add_heading_styled('7.1 問題描述', level=2)
add_para('題庫中的第 11、12、13 題為連貫題組，三題共用同一棵樹的圖示，分別從不同角度進行判斷：')

add_table(
    ['題號', '題目', '說明'],
    [
        ['11', '給定下列的樹，請問這棵樹是不是「完整二元樹」？', '基礎判斷'],
        ['12', '承上題，請問這棵樹是不是「完滿二元樹」？', '延伸判斷'],
        ['13', '承上題，請問這棵樹是不是「二元搜尋樹」？', '延伸判斷'],
    ],
    col_widths=[2, 10, 3]
)

add_para('若系統隨機抽題時只抽到其中一題，學生將無法回答該題，因為缺少前後文的參照。因此必須實作保護機制，確保這三題要麼一起出現、要麼都不出現。')

add_heading_styled('7.2 解決方案', level=2)
add_para('在 views.py 中實作 _sample_with_groups() 函數，演算法如下：')

steps = [
    '將所有題目依照 sequence_group 欄位分組（有設定值的為一組，沒有的各自一題）。',
    '同組內的題目按題號排序（確保 11→12→13 的順序）。',
    '將各組打成一個群組列表，並打亂群組間的順序。',
    '從頭開始依序將整組題目加入結果，直到達到目標題數為止。',
]
for i, s in enumerate(steps, 1):
    add_para(f'{i}. {s}')

add_para('')
add_para('不同選題數的結果：', bold=True)
add_table(
    ['選題數', '結果'],
    [
        ['5 題', '隨機抽取，若抽到 11→12→13 則整組 3 題都進入'],
        ['10 題', '同上原則'],
        ['15 題', '同上原則'],
        ['20 題', '同上原則'],
        ['23 題（全部）', '全部題目隨機排列，11→12→13 保持順序'],
    ],
    col_widths=[3, 12]
)

# ══════════════════════════════════════════
# 八、間隔學習演算法
# ══════════════════════════════════════════
add_heading_styled('八、間隔學習演算法', level=1)

add_heading_styled('8.1 演算法流程', level=2)
steps = [
    '第一輪：正常作答所有題目，記錄答對題數作為最終評分依據。',
    '若全部答對，直接結算存檔。',
    '若有錯題，將錯題挑出，進入間隔學習輪次。',
    '間隔學習輪次中，只出現尚未答對的題目。',
    '答對的題目從錯題清單中移除，下一輪不再出現。',
    '持續進行，直到錯題清單為空（全部答對）為止，無輪次上限。',
]
for i, s in enumerate(steps, 1):
    add_para(f'{i}. {s}')

add_para('')
add_heading_styled('8.2 評分方式', level=2)
add_para('最終分數 =（第一輪答對數 ÷ 總題數）× 100', bold=True)
add_para('後續間隔學習輪次只為了鞏固學習，不影響分數計算。此設計確保分數忠實反映學生的初始理解程度，同時透過反覆練習強化弱點。')

add_heading_styled('8.3 排行榜排名依據', level=2)
add_table(
    ['排行類型', '排名依據', '篩選條件'],
    [
        ['一般排行', '分數由高至低', '章節、題數'],
        ['間隔學習排行', '完成時間由短至長', '章節、題數'],
    ],
    col_widths=[4, 4, 4]
)

# ══════════════════════════════════════════
# 九、開發時程
# ══════════════════════════════════════════
add_heading_styled('九、開發時程', level=1)

add_table(
    ['階段', '項目', '預估工時', '備註'],
    [
        ['1', '環境建置與 Django 專案初始化', '2 小時', 'Python、Django、SQLite'],
        ['2', '資料庫模型設計（Model）', '3 小時', 'User / Question / QuizRecord / WrongAnswer'],
        ['3', '題目匯入腳本開發', '3 小時', '從 Word 文件匯入 23 題'],
        ['4', '用戶系統（註冊／登入／登出）', '4 小時', '含個人資料管理'],
        ['5', '測驗核心邏輯（出題／作答／計分）', '6 小時', '含選項隨機排列'],
        ['6', '間隔學習演算法實作', '4 小時', 'SR 流程與 Session 管理'],
        ['7', '結果頁面與作答記錄', '3 小時', '成績結算、錯題記錄'],
        ['8', '錯題複習功能', '3 小時', '錯題列表、詳情、複習'],
        ['9', '排行榜系統', '3 小時', '一般排行 + SR 排行'],
        ['10', '管理後台', '6 小時', '題目／用戶／記錄 CRUD'],
        ['11', '隨機挑戰功能', '3 小時', '跨章節混合出題設定'],
        ['12', 'Aurora Glass 主題 CSS 設計', '5 小時', '玻璃擬態、漸層、動畫'],
        ['13', '首頁動畫效果', '2 小時', '流星雨、游標光暈'],
        ['14', '連貫題保護機制', '2 小時', '_sample_with_groups()'],
        ['15', '管理後台美化', '2 小時', 'Django Admin 自訂樣式'],
        ['16', '測試與除錯', '4 小時', '邊際情況、相容性測試'],
        ['17', '文件撰寫', '3 小時', 'README、計畫書'],
        ['', '合計', '58 小時', ''],
    ],
    col_widths=[2, 6, 2.5, 5]
)

# ══════════════════════════════════════════
# 十、資料庫操作說明
# ══════════════════════════════════════════
add_heading_styled('十、資料庫操作說明', level=1)

add_para('10.1 建立管理員帳號', bold=True)
add_para('python manage.py createsuperuser')

add_para('10.2 匯入題庫', bold=True)
add_para('python manage.py import_questions')

add_para('10.3 備份資料庫', bold=True)
add_para('copy db.sqlite3 db.sqlite3.backup')

add_para('10.4 重置資料庫', bold=True)
add_para('del db.sqlite3')
add_para('rmdir /s migrations')
add_para('python manage.py makemigrations quiz_app')
add_para('python manage.py migrate')

# ══════════════════════════════════════════
# 十一、URL 路由一覽
# ══════════════════════════════════════════
add_heading_styled('十一、URL 路由一覽', level=1)

add_table(
    ['路徑', '檢視函數', '說明'],
    [
        ['/', 'quiz_home', '首頁'],
        ['/login/', 'login_view', '登入'],
        ['/register/', 'register_view', '註冊'],
        ['/logout/', 'logout_view', '登出'],
        ['/quiz/{chapter}/', 'start_quiz', '測驗設定頁'],
        ['/take-quiz/{chapter}/', 'take_quiz', '開始作答'],
        ['/take-random-quiz/', 'take_random_quiz', '開始隨機測驗'],
        ['/random-quiz-setup/', 'random_quiz_setup', '隨機測驗設定'],
        ['/submit-quiz/', 'submit_quiz', '提交答案'],
        ['/leaderboard/', 'leaderboard', '一般排行'],
        ['/leaderboard/sr/', 'sr_leaderboard', 'SR 排行'],
        ['/records/', 'quiz_records', '作答記錄'],
        ['/wrong-answers/', 'wrong_answers', '錯題列表'],
        ['/wrong-answers/{id}/', 'wrong_answers_detail', '錯題詳情'],
        ['/review-wrong/{id}/', 'review_wrong', '錯題複習'],
        ['/admin-panel/', 'admin_panel', '自訂管理後台'],
    ],
    col_widths=[5, 4, 6]
)

# ══════════════════════════════════════════
# 十二、未來展望
# ══════════════════════════════════════════
add_heading_styled('十二、未來展望', level=1)

future_items = [
    ('題庫擴充', '新增更多章節（如排序、圖形、雜湊表），擴大全域題庫，涵蓋資料結構課程全部重點。'),
    ('多題型支援', '在現有選擇題基礎上，加入是非題、填空題、程式碼實作題等多元題型。'),
    ('圖表分析', '以視覺化圖表呈現學習趨勢、弱點分析、進步曲線，幫助學生掌握學習狀況。'),
    ('匯出功能', '支援將錯題、成績匯出為 PDF 或 Excel 格式，方便學生離線複習。'),
    ('批次匯入改善', '支援 CSV / JSON 格式批次匯入題目，降低題庫維護成本。'),
    ('API 開放', '提供 RESTful API 供第三方應用程式串接，擴展系統應用場景。'),
    ('行動裝置支援', '以 PWA 技術或 React Native 封裝成行動應用，實現跨平台學習體驗。'),
]

for title, desc in future_items:
    p = doc.add_paragraph()
    run_b = p.add_run(f'• {title}：')
    run_b.bold = True
    run_b.font.name = 'Times New Roman'
    run_b.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    run_n = p.add_run(desc)
    run_n.font.name = 'Times New Roman'
    run_n.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

# ══════════════════════════════════════════
# 十三、參考資料
# ══════════════════════════════════════════
add_heading_styled('十三、參考資料', level=1)

refs = [
    'Django 6.0 官方文件 — https://docs.djangoproject.com/',
    '資料結構 — 樹狀結構（第 7 章課程教材）',
    'CSS Glassmorphism — https://glassmorphism.com/',
    'Spaced Repetition（間隔學習）— https://en.wikipedia.org/wiki/Spaced_repetition',
    'MDN Web Docs — https://developer.mozilla.org/',
    'python-docx 文件 — https://python-docx.readthedocs.io/',
]
for i, r in enumerate(refs, 1):
    add_para(f'{i}. {r}')

# ── 儲存 ──
output_path = os.path.join(os.path.dirname(__file__), '專題計畫書.docx')
doc.save(output_path)
print(f'已產生：{output_path}')